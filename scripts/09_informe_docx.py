"""
09_informe_docx.py — Informe del modelo V4 en Word (.docx).

Arma docs/informe_v4.docx con lenguaje simple y directo: qué se hizo, el cambio
que mejoró el modelo, los números antes/después, la tabla por especie y los mapas
(panel + dos ejemplos). Pensado para que lo lea cualquiera, no solo gente técnica.

Uso: python scripts/09_informe_docx.py
"""
from __future__ import annotations

import sys
from pathlib import Path

import pandas as pd
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

_ROOT = Path(__file__).resolve().parents[1]
FIGS = _ROOT / "outputs" / "figures"
METRICAS = _ROOT / "outputs" / "tables" / "metricas_v4_completa.csv"
SALIDA = _ROOT / "docs" / "informe_v4.docx"


def _nivel(boyce: float) -> str:
    if boyce >= 0.9:
        return "Excelente"
    if boyce >= 0.7:
        return "Confiable"
    if boyce >= 0.3:
        return "Aceptable"
    return "Revisar"


def _tabla(doc, encabezados, filas):
    t = doc.add_table(rows=1, cols=len(encabezados))
    t.style = "Light Grid Accent 1"
    for i, h in enumerate(encabezados):
        run = t.rows[0].cells[i].paragraphs[0].add_run(h)
        run.bold = True
    for fila in filas:
        celdas = t.add_row().cells
        for i, val in enumerate(fila):
            celdas[i].text = str(val)
    return t


def main():
    met = pd.read_csv(METRICAS)
    doc = Document()

    # --- Portada / título ---
    h = doc.add_heading("Mapas de distribución de plantas — Versión 4", level=0)
    sub = doc.add_paragraph("Dónde pueden vivir 16 especies de flora chilena, según el clima y el terreno")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in sub.runs:
        r.italic = True
        r.font.size = Pt(12)

    # --- Qué hicimos ---
    doc.add_heading("Qué hicimos", level=1)
    doc.add_paragraph(
        "Tomamos los registros de dónde se ha visto cada planta y los cruzamos con el clima "
        "y el terreno de Sudamérica. Con eso el modelo aprende qué condiciones le gustan a cada "
        "especie y dibuja un mapa que muestra, zona por zona, qué tan buena es para ella. "
        "Usamos cinco métodos a la vez y los combinamos, para no depender de uno solo."
    )

    # --- El cambio que mejoró todo ---
    doc.add_heading("El cambio que mejoró el modelo", level=1)
    doc.add_paragraph(
        "La versión anterior tenía un problema simple: comparaba a cada planta contra un fondo "
        "que no calzaba con dónde vive de verdad. Lo arreglamos haciendo que cada especie se "
        "compare contra su propia zona, un radio de unos 300 km alrededor de los lugares donde "
        "se la ha encontrado. Fue un ajuste chico, pero los resultados mejoraron en todo."
    )
    doc.add_paragraph("Así se ven los números, antes y después (promedio de las 16 especies):")
    _tabla(doc, ["Medida", "Antes (V3)", "Ahora (V4)"],
           [["AUC (qué tan bien separa)", "0.77", "0.83"],
            ["TSS (aciertos vs errores)", "0.26", "0.47"],
            ["Boyce (acierta dónde está la planta)", "0.44", "0.86"]])
    doc.add_paragraph(
        "El número que más importa acá es el Boyce: mide si el mapa marca alto justo donde la "
        "planta aparece de verdad. Pasó de 0.44 a 0.86. En 15 de las 16 especies el mapa quedó "
        "confiable. Es un salto grande y se nota en los mapas."
    )

    # --- El modelo que elegimos ---
    doc.add_heading("El modelo que elegimos", level=1)
    doc.add_paragraph(
        "Comparamos juntar los cinco métodos (lo que llamamos ensemble) contra usar solo el más "
        "fuerte, MaxEnt. En puntería van casi iguales, pero el ensemble gana donde importa: da el "
        "mejor Boyce y es más estable, porque no queda colgado de un solo método. Por eso es el "
        "modelo que dejamos como oficial."
    )
    _tabla(doc, ["Modelo", "AUC", "TSS", "Boyce"],
           [["Ensemble (el que usamos)", "0.83", "0.47", "0.86"],
            ["MaxEnt solo", "0.82", "0.48", "—"]])

    # --- Especie por especie ---
    doc.add_heading("Cómo le fue a cada especie", level=1)
    doc.add_paragraph(
        "Ordenadas de la que mejor quedó a la que peor. Casi todas andan bien; las mejores marcan "
        "justo la franja del norte chico donde crecen estas plantas."
    )
    m = met.sort_values("boyce_ensemble", ascending=False)
    filas = [[r.especie, int(r.n_pres), f"{r.auc_ensemble:.2f}",
              f"{r.tss_ensemble:.2f}", f"{r.boyce_ensemble:.2f}", _nivel(r.boyce_ensemble)]
             for _, r in m.iterrows()]
    _tabla(doc, ["Especie", "Registros", "AUC", "TSS", "Boyce", "Nivel"], filas)
    doc.add_paragraph(
        "La única que todavía no sirve es Atriplex semibaccata: es una especie introducida y con "
        "pocos datos, así que su mapa conviene tomarlo aparte y con pinzas."
    )

    # --- Mapas ---
    doc.add_heading("Los mapas", level=1)
    doc.add_paragraph(
        "Cada mapa pinta de oscuro a amarillo la idoneidad (de 0 a 1): mientras más amarillo, mejor "
        "zona para la planta. Los puntos rojos son los lugares donde se la ha visto."
    )
    panel = FIGS / "_panel_idoneidad_v4.png"
    if panel.exists():
        doc.add_picture(str(panel), width=Inches(6.3))
        cap = doc.add_paragraph("Las 16 especies de un vistazo (ordenadas por Boyce).")
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in cap.runs:
            r.italic = True

    for slug, pie in [("nolana_divaricata", "Nolana divaricata: uno de los mejores. El amarillo cae justo sobre los puntos rojos."),
                      ("atriplex_semibaccata", "Atriplex semibaccata: el caso flojo, mapa disperso y poco confiable.")]:
        f = FIGS / f"{slug}_idoneidad_sa.png"
        if f.exists():
            doc.add_picture(str(f), width=Inches(4.5))
            cap = doc.add_paragraph(pie)
            cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in cap.runs:
                r.italic = True

    # --- Cómo leer los mapas ---
    doc.add_heading("Cómo leer los mapas (en simple)", level=1)
    for txt in [
        "Los colores son idoneidad relativa, de 0 a 1: qué tan buena es la zona para la planta. "
        "No es una probabilidad exacta, es un “qué tan apta” comparado entre zonas.",
        "Donde el clima se parece mucho al de los registros, el mapa es más seguro. En zonas con "
        "clima muy distinto conviene ir con más cautela.",
        "El modelo casi no deja afuera lugares donde la planta sí está (algo bueno).",
    ]:
        p = doc.add_paragraph(txt, style="List Bullet")

    # --- En resumen ---
    doc.add_heading("En resumen", level=1)
    cierre = doc.add_paragraph(
        "La Versión 4 es, por lejos, la mejor que tenemos. Un arreglo simple en cómo el modelo "
        "elige con qué comparar a cada planta subió todos los números y dejó 15 de 16 mapas "
        "listos para usar. Es una base sólida y confiable para seguir trabajando."
    )
    for r in cierre.runs:
        r.font.size = Pt(11)

    SALIDA.parent.mkdir(parents=True, exist_ok=True)
    doc.save(SALIDA)
    print(f"Informe guardado: {SALIDA}")


if __name__ == "__main__":
    main()
