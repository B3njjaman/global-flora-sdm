"""
10_md_a_docx.py — Convierte los documentos Markdown de docs/ a Word (.docx).

Conversor propio (no requiere pandoc) basado en python-docx. Soporta lo que usan
los documentos del proyecto: títulos (#..######), párrafos, listas con viñetas y
numeradas, tablas, citas (>), bloques de código (```), y formato inline
(**negrita**, *cursiva*, `código`, [texto](enlace)).

Convierte cada docs/**/*.md a un .docx hermano (mismo nombre). Salta informe_v4
(ya existe como .docx propio).

Uso: python scripts/10_md_a_docx.py
"""
from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor

_ROOT = Path(__file__).resolve().parents[1]
DOCS = _ROOT / "docs"

_LINK = re.compile(r"\[([^\]]+)\]\([^)]+\)")
_INLINE = re.compile(r"(\*\*.+?\*\*|`[^`]+`|\*[^*]+\*)")
_SEP = re.compile(r"^\s*\|?[\s:|-]*-[\s:|-]*\|?\s*$")  # fila separadora de tabla


def _limpiar(texto: str) -> str:
    """Reemplaza [texto](url) por el texto visible."""
    return _LINK.sub(r"\1", texto)


def _add_runs(parrafo, texto: str) -> None:
    """Agrega 'texto' a un párrafo respetando **negrita**, *cursiva* y `código`."""
    texto = _limpiar(texto)
    for tok in _INLINE.split(texto):
        if not tok:
            continue
        if tok.startswith("**") and tok.endswith("**") and len(tok) > 4:
            parrafo.add_run(tok[2:-2]).bold = True
        elif tok.startswith("`") and tok.endswith("`") and len(tok) > 2:
            run = parrafo.add_run(tok[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(9.5)
        elif tok.startswith("*") and tok.endswith("*") and len(tok) > 2:
            parrafo.add_run(tok[1:-1]).italic = True
        else:
            parrafo.add_run(tok)


def _celdas(linea: str) -> list[str]:
    s = linea.strip()
    if s.startswith("|"):
        s = s[1:]
    if s.endswith("|"):
        s = s[:-1]
    return [c.strip() for c in s.split("|")]


def _es_fila_tabla(linea: str) -> bool:
    return "|" in linea and linea.strip().startswith(("|", "")) and linea.count("|") >= 1


def convertir(md_path: Path, docx_path: Path) -> None:
    lineas = md_path.read_text(encoding="utf-8").splitlines()
    doc = Document()
    i, n = 0, len(lineas)
    parrafo_buf: list[str] = []

    def flush_parrafo():
        nonlocal parrafo_buf
        if parrafo_buf:
            p = doc.add_paragraph()
            _add_runs(p, " ".join(parrafo_buf).strip())
            parrafo_buf = []

    while i < n:
        ln = lineas[i]
        s = ln.strip()

        # --- bloque de código ---
        if s.startswith("```"):
            flush_parrafo()
            i += 1
            code: list[str] = []
            while i < n and not lineas[i].strip().startswith("```"):
                code.append(lineas[i])
                i += 1
            i += 1  # cierre ```
            p = doc.add_paragraph()
            for j, c in enumerate(code):
                if j:
                    p.add_run().add_break()
                run = p.add_run(c)
                run.font.name = "Consolas"
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
            continue

        # --- línea en blanco ---
        if not s:
            flush_parrafo()
            i += 1
            continue

        # --- título ---
        if s.startswith("#"):
            flush_parrafo()
            nivel = len(s) - len(s.lstrip("#"))
            texto = s[nivel:].strip()
            doc.add_heading(_limpiar(texto), level=min(nivel, 9))
            i += 1
            continue

        # --- regla horizontal ---
        if re.fullmatch(r"-{3,}|\*{3,}|_{3,}", s):
            flush_parrafo()
            i += 1
            continue

        # --- tabla (fila + separadora debajo) ---
        if "|" in ln and i + 1 < n and _SEP.match(lineas[i + 1]):
            flush_parrafo()
            encabezados = _celdas(ln)
            i += 2  # salta encabezado + separadora
            filas = []
            while i < n and "|" in lineas[i] and lineas[i].strip():
                filas.append(_celdas(lineas[i]))
                i += 1
            ncol = len(encabezados)
            t = doc.add_table(rows=1, cols=ncol)
            try:
                t.style = "Light Grid Accent 1"
            except Exception:
                t.style = "Table Grid"
            for c, h in enumerate(encabezados):
                cell = t.rows[0].cells[c]
                cell.paragraphs[0].text = ""
                run = cell.paragraphs[0].add_run(_limpiar(h))
                run.bold = True
            for fila in filas:
                celdas = t.add_row().cells
                for c in range(ncol):
                    val = fila[c] if c < len(fila) else ""
                    celdas[c].paragraphs[0].text = ""
                    _add_runs(celdas[c].paragraphs[0], val)
            continue

        # --- cita ---
        if s.startswith(">"):
            flush_parrafo()
            texto = s.lstrip(">").strip()
            p = doc.add_paragraph(style="Intense Quote") if "Intense Quote" in [st.name for st in doc.styles] else doc.add_paragraph()
            _add_runs(p, texto)
            i += 1
            continue

        # --- lista con viñetas ---
        m = re.match(r"^\s*[-*+]\s+(.*)$", ln)
        if m:
            flush_parrafo()
            p = doc.add_paragraph(style="List Bullet")
            _add_runs(p, m.group(1))
            i += 1
            continue

        # --- lista numerada ---
        m = re.match(r"^\s*\d+\.\s+(.*)$", ln)
        if m:
            flush_parrafo()
            p = doc.add_paragraph(style="List Number")
            _add_runs(p, m.group(1))
            i += 1
            continue

        # --- texto normal (se acumula en párrafo) ---
        parrafo_buf.append(s)
        i += 1

    flush_parrafo()
    docx_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(docx_path)


def main():
    objetivo = sorted(DOCS.rglob("*.md"))
    if not objetivo:
        print("No hay .md en docs/.")
        return
    print(f"Convirtiendo {len(objetivo)} documentos a .docx...")
    for md in objetivo:
        docx = md.with_suffix(".docx")
        try:
            convertir(md, docx)
            print(f"  OK {md.relative_to(_ROOT)} -> {docx.name}")
        except Exception as exc:  # noqa: BLE001
            print(f"  ERROR {md.name}: {exc}")
    print(f"Listo. .docx en {DOCS}/ (y docs/v4/).")


if __name__ == "__main__":
    main()
